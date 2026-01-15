from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple, Optional

from src.core.exceptions import ProcessingError
from src.core.models.office import OfficeTextContent
from ._common import sanitize_filename, ensure_dir


@dataclass(slots=True)
class DocxExtraction:
    markdown: str
    image_paths: List[str]
    text_contents: List[OfficeTextContent]


def _heading_prefix(style_name: str) -> Optional[str]:
    # DOCX Styles: "Heading 1", "Heading 2", ...
    if not style_name:
        return None
    s = style_name.strip().lower()
    if not s.startswith("heading"):
        return None
    parts = s.split()
    if len(parts) != 2:
        return None
    try:
        level = int(parts[1])
    except Exception:
        return None
    level = max(1, min(level, 6))
    return "#" * level


def extract_docx_to_markdown(input_path: Path, images_dir: Path) -> DocxExtraction:
    """Extrahiert DOCX nach Markdown + Embedded-Images.

    Implementation bewusst „einfach und robust“:
    - Headings werden aus dem Paragraph-Style abgeleitet
    - Inline-Images werden über die Run-XML (`a:blip @r:embed`) erkannt
    - Tabellen werden als einfache Markdown-Tabellen ausgegeben
    """
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise ProcessingError(f"python-docx ist nicht installiert: {e}")

    ensure_dir(images_dir)
    doc = Document(str(input_path))

    # Bild-Relationship ID → gespeicherter Dateiname
    rid_to_filename: dict[str, str] = {}
    image_paths: List[str] = []
    text_contents: List[OfficeTextContent] = []
    lines: List[str] = []

    # Namespace-Map für XPath (python-docx nutzt lxml)
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    def _save_image_rid(rid: str) -> Optional[str]:
        if rid in rid_to_filename:
            return rid_to_filename[rid]
        try:
            part = doc.part.related_parts[rid]
            blob: bytes = part.blob  # type: ignore[attr-defined]
            # partname kann z.B. '/word/media/image1.png' sein
            guessed = getattr(part, "partname", None)
            stem = f"docx-img-{len(rid_to_filename) + 1}"
            suffix = ".bin"
            if guessed:
                name = str(guessed).split("/")[-1]
                if "." in name:
                    suffix = "." + name.split(".")[-1].lower()
            filename = sanitize_filename(stem) + suffix
            out_path = images_dir / filename
            with open(out_path, "wb") as f:
                f.write(blob)
            rel_path = f"images/{filename}"
            rid_to_filename[rid] = rel_path
            image_paths.append(rel_path)
            return rel_path
        except Exception:
            return None

    # Paragraphs + Tables in Dokument-Reihenfolge (docx hat keine einheitliche API)
    # Für Einfachheit: erst Paragraphs, dann Tables. Das ist nicht layout-treu,
    # aber stabil. (Pipeline B ist für layout-treue Vergleiche da.)
    block_index = 0
    for p in doc.paragraphs:
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        prefix = _heading_prefix(style_name)

        # Text + Inline-Images in Run-Reihenfolge
        parts: List[str] = []
        for run in p.runs:
            # Text
            if run.text:
                parts.append(run.text)
            # Inline images via blip embed rid
            try:
                blips = run.element.xpath(".//a:blip", namespaces=ns)  # type: ignore[attr-defined]
                for blip in blips:
                    rid = blip.get(f"{{{ns['r']}}}embed")
                    if rid:
                        rel = _save_image_rid(str(rid))
                        if rel:
                            parts.append(f" ![]({rel}) ")
            except Exception:
                # XML Parsing-Fehler ignorieren – Text bleibt erhalten
                pass

        raw = "".join(parts).strip()
        if not raw:
            continue
        if prefix:
            line = f"{prefix} {raw}"
        else:
            line = raw
        lines.append(line)
        text_contents.append(OfficeTextContent(index=block_index, text=raw, kind="docx_paragraph"))
        block_index += 1

    # Tabellen: einfache Markdown-Tabelle (erste Zeile als Header)
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        if not rows:
            continue
        # Normalize column count
        col_count = max(len(r) for r in rows)
        norm = [r + [""] * (col_count - len(r)) for r in rows]
        header = norm[0]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for body_row in norm[1:]:
            lines.append("| " + " | ".join(body_row) + " |")
        table_text = "\n".join(["\t".join(r) for r in norm])
        text_contents.append(OfficeTextContent(index=block_index, text=table_text, kind="docx_table"))
        block_index += 1

    markdown = "\n\n".join(lines).strip() + "\n"
    return DocxExtraction(markdown=markdown, image_paths=image_paths, text_contents=text_contents)






