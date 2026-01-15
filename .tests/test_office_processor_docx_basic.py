from __future__ import annotations

import asyncio
import tempfile
from pathlib import Path

import pytest

from src.processors.office_processor import OfficeProcessor


def _make_docx(tmp: Path) -> Path:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx nicht verfügbar: {e}")

    p = tmp / "sample.docx"
    doc = Document()
    doc.add_heading("Titel", level=1)
    doc.add_paragraph("Dies ist ein Test.")
    # Tabelle
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"
    doc.save(str(p))
    return p


def test_office_processor_docx_creates_markdown() -> None:
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        docx_path = _make_docx(tmp)

        proc = OfficeProcessor(process_id="test-docx")
        res = asyncio.run(
            proc.process(
                docx_path,
                include_images=True,
                include_previews=True,
                use_cache=False,
                base_cache_dir=tmp / "cache",
            )
        )

        assert res.data.metadata.format == "docx"
        assert res.data.extracted_text
        assert "Titel" in res.data.extracted_text
        assert "Dies ist ein Test." in res.data.extracted_text
        assert (Path(res.markdown_path)).exists()






